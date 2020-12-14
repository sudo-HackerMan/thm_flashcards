#!/usr/bin/python3
#-*- coding: utf-8 -*-

import sys
import os
from PyQt5.QtWidgets import QApplication, QWidget, QFileDialog, QMessageBox
from PyQt5 import uic
import time
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtGui import QIcon
from threading import Thread
import shutil
import zipfile
import xlwt
from docx import Document
from docx.shared import Mm, Pt


os.environ['QT_STYLE_OVERRIDE'] = 'fusion'


class setProgressBarValueClass(QThread):
    def __init__(self, gameWindow, parent=None):
        super().__init__()
        self.gameWindow = gameWindow

    def run(self):
        while self.gameWindow.showResultLabel.text() != 'END':
            self.gameWindow.myProgressBar.setValue(int(self.gameWindow.progressLabel2.text()))
            time.sleep(0.7)
        else:
            self.gameWindow.myProgressBar.setValue(100)


class gameThreadClass(QThread):
    def __init__(self, gameWindow, parent=None):
        super().__init__()
        self.gameWindow = gameWindow


    def samePercent(self, answerC, answerG):
        myBuffer = answerG
        result = 0
        for sym in answerC:
            if sym in myBuffer:
                myBuffer = myBuffer.replace(sym, '', 1)
                result += 1
        r = int(result / len(answerC) * 100)
        _r = int(result / len(answerG) * 100)
        return int('{}'.format(r if r < _r else _r))


    def showResult(self, correctCount, allcards):
        endResult = 'Последний результат: ' + str(correctCount) + '/' + str(allcards)
        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)
        msgBox.setText(endResult)
        msgBox.setWindowTitle('Информация')
        msgBox.setStandardButtons(QMessageBox.Ok)
        msgBox.exec()
        sys.exit()


    def run(self):
        cN,cB,cAA,cK,cDK = False,False,False,False,False
        count = 0
        correctCount = 0
        currentStack = self.gameWindow.stackNameLabel.text()
        self.currentCard = ''
        allcards = 0

        for k in os.listdir('user_data/' + currentStack):
            if k[-5:] == '.card':
                allcards += 1

        self.gameWindow.countProgressLabel.setText(str(count) + '/' + str(allcards))
        for j in os.listdir('user_data/' + currentStack):
            if j[-5:] == '.card':
                self.currentCard = j
                self.gameWindow.currentCardLabel.setText(j)
                myCard = open('user_data/' + currentStack + '/' + j, 'r')
                self.gameWindow.cardNameLabel.setText(myCard.readline()[:-1])
                description = myCard.readline()[:-1]
                self.gameWindow.cardDescriptionTextEdit.clear()
                self.gameWindow.cardDescriptionTextEdit.insertPlainText(description)
                self.gameWindow.update()
                answer = myCard.readline()[:-1]
                picture = myCard.readline()
                pixmap = QPixmap(os.getcwd() + '/' + picture)
                self.gameWindow.pictureBox.setPixmap(pixmap)
                myCard.close()


                while not(cN == True or cB == True or cAA == True or cK == True or cDK == True):
                    cN = bool(self.gameWindow.nextBoolLabel.text())
                    cB = bool(self.gameWindow.backBoolLabel.text())
                    cAA = bool(self.gameWindow.acceptBoolLabel.text())
                    cK = bool(self.gameWindow.knowBoolLabel.text())
                    cDK = bool(self.gameWindow.dontKnowBoolLabel.text())
                    time.sleep(0.3)
                else:
                    if cAA:
                        #True if self.samePercent(answerFromCard, givenAnswer) >= 50 else False
                        if self.gameWindow.answerBox.text() != '':
                            if self.samePercent(answer, self.gameWindow.answerBox.text()) >= 50:
                                self.gameWindow.showResultLabel.setText('OK')
                                self.gameWindow.showResultLabel.setStyleSheet('color: green;')
                                correctCount += 1
                                self.gameWindow.correctCount.setText(str(int(self.gameWindow.correctCount.text()) + 1))

                            else:
                                self.gameWindow.showResultLabel.setText('BAD')
                                self.gameWindow.showResultLabel.setStyleSheet('color: red;')
                        else:
                            self.gameWindow.showResultLabel.setText('BAD')
                            self.gameWindow.showResultLabel.setStyleSheet('color: red;')

                    elif cK:
                        correctCount += 1
                        self.gameWindow.correctCount.setText(str(int(self.gameWindow.correctCount.text()) + 1))
                        self.gameWindow.showResultLabel.setText('OK')
                        self.gameWindow.showResultLabel.setStyleSheet('color: green;')


                    elif cDK:
                        self.gameWindow.showResultLabel.setText('BAD')
                        self.gameWindow.showResultLabel.setStyleSheet('color: red;')

                    count += 1
                    self.gameWindow.nextBoolLabel.setText('')
                    self.gameWindow.backBoolLabel.setText('')
                    self.gameWindow.acceptBoolLabel.setText('')
                    self.gameWindow.knowBoolLabel.setText('')
                    self.gameWindow.dontKnowBoolLabel.setText('')

                    cN,cB,cAA,cK,cDK = False,False,False,False,False

                self.gameWindow.countProgressLabel.setText(str(count) + '/' + str(allcards))
                currentProgress = 100 * count // allcards
                self.gameWindow.progressLabel2.setText(str(currentProgress))
                self.gameWindow.answerBox.clear()

        self.gameWindow.showResultLabel.setText('END')
        self.gameWindow.showResultLabel.setStyleSheet('color: blue;')
        result = str(correctCount) + '/' + str(allcards)

        self.gameWindow.showResultLabel.setText('OK/BAD/END')
        self.gameWindow.showResultLabel.setStyleSheet('')
        self.showResult(str(correctCount), str(allcards))
        sys.exit()

        count = 0
        correctCount = 0


class App(QWidget):
    def __init__(self):
        self.start()
        self.set()


    def start(self):
        self.ui = uic.loadUi('main_form.ui')
        self.aboutWindow = uic.loadUi('aboutSW.ui')
        self.selectStack = uic.loadUi('selectStack.ui')
        self.gameWindow = uic.loadUi('game_window.ui')
        self.editor = uic.loadUi('editor.ui')
        self.createStack = uic.loadUi('create_stack.ui')
        self.importStack = uic.loadUi('stackImport.ui')
        #self.gameOver = uic.loadUi('end.ui')
        self.exportStack = uic.loadUi('export.ui')
        self.editor.descTextEditBox.setToolTip('''Описание карточки - это то, что будет отображаться
под вопросом карточки. Туда можно написать, что
конкретно нужно ответить.
Пример: Передняя сторона: Великая Отечественная Война,
Описание: Годы?, Ответ: 1941-1945.''')
        self.editor.descCaptLabel.setToolTip('''Описание карточки - это то, что будет отображаться
под вопросом карточки. Туда можно написать, что
конкретно нужно ответить.
Пример: Передняя сторона: Великая Отечественная Война,
Описание: Годы?, Ответ: 1941-1945.''')
        self.editor.fileNameLabel.setToolTip('''Имя файла карточки - это название файла, которое
будет использоваться для обращения к карточке во время
игры. Может содержать только латинские буквы.''')
        self.editor.cardFileNameBox.setToolTip('''Имя файла карточки - это название файла, которое
будет использоваться для обращения к карточке во время
игры. Может содержать только латинские буквы.''')
        self.ui.show()


    def importArchive(self):
        stackToExportName = ''
        if bool(self.importStack.pathQLineEdit.text()):

            archivePath = self.importStack.pathQLineEdit.text()
            if zipfile.is_zipfile(archivePath):
                zF_ud = zipfile.ZipFile(archivePath, 'r')
                zF_ud.extract('stackName', os.getcwd() + '/tmp/')
                zF_ud.close()

                stackToExportName = open(os.getcwd() + '/tmp/stackName').readline()[:-1]
                if not(os.path.exists('user_data/' + stackToExportName)):
                    zF_u = zipfile.ZipFile(archivePath, 'r')
                    zF_u.extractall('user_data/' + stackToExportName)
                    zF_u.close()
                    os.remove(os.getcwd() + '/tmp/stackName')

                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText('Стопка карт была успешно импортирована. Перезапустите программу для обновления информации.')
                    msgBox.setWindowTitle('Информация')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()
                    sys.exit()
                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Critical)
                    msgBox.setText('Стопка с таким именем уже существует.')
                    msgBox.setWindowTitle('Ошибка')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()
            else:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Critical)
                msgBox.setText('Указанный файл не является *.zip архивом.')
                msgBox.setWindowTitle('Ошибка')
                msgBox.setStandardButtons(QMessageBox.Ok)
                msgBox.exec()
        else:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Critical)
            msgBox.setText('Следует указать архив для импорта.')
            msgBox.setWindowTitle('Ошибка')
            msgBox.setStandardButtons(QMessageBox.Ok)
            msgBox.exec()

    def exportStackToXLS(self, stackName, path):
        myStyle = xlwt.easyxf('font: name Times New Roman, color-index black, bold off')
        wb = xlwt.Workbook()
        ws = wb.add_sheet(stackName)
        ws.write(0, 0, 'Передняя сторона', xlwt.easyxf('font: name Times New Roman, color-index black, bold on'))
        ws.write(0, 1, 'Описание', xlwt.easyxf('font: name Times New Roman, color-index black, bold on'))
        ws.write(0, 2, 'Ответ', xlwt.easyxf('font: name Times New Roman, color-index black, bold on'))

        for card in range(0, len(os.listdir('user_data/' + stackName))):
            fsList = os.listdir('user_data/' + stackName)
            if fsList[card][-5:] == '.card':
                cardReader = open('user_data/' + stackName + '/' + fsList[card])
                frontSide = cardReader.readline()[:-1]
                description = cardReader.readline()[:-1]
                backSide = cardReader.readline()[:-1]
                ws.write(card + 1, 0, frontSide, myStyle)
                ws.write(card + 1, 1, description, myStyle)
                ws.write(card + 1, 2, backSide, myStyle)

        wb.save(path)

        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)
        msgBox.setText('Стопка была успешно экспортирована в ' + path)
        msgBox.setWindowTitle('Информация')
        msgBox.setStandardButtons(QMessageBox.Ok)
        msgBox.exec()


    def saveArchive(self):
        currentStack = self.exportStack.stackLabel.text()
        _myFileName = QFileDialog.getSaveFileName()
        _path = _myFileName[0]
        if self.exportStack.actionListWidget.currentItem().text() == 'Экспорт стопки в архив (для копирования)':
            if bool(_path):
                if not(os.path.exists(_path) or os.path.exists(_path + '.zip')):
                        _path = _path if _path[-4:] == '.zip' else _path + '.zip'
                        zF = zipfile.ZipFile(_path, 'w')
                        os.chdir(currentStack)
                        for file in os.listdir(os.getcwd()):
                            zF.write(file)
                        open('stackName', 'w').write(currentStack[10:] + '\n')
                        zF.write('stackName')
                        zF.close()
                        os.remove('stackName')
                        os.chdir(__file__[:-8])
                        msgBox = QMessageBox()
                        msgBox.setIcon(QMessageBox.Information)
                        msgBox.setText('Стопка была успешно экспортирована в ' + _path)
                        msgBox.setWindowTitle('Информация')
                        msgBox.setStandardButtons(QMessageBox.Ok)
                        msgBox.exec()

                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Critical)
                    msgBox.setText('Архив с таким названием уже существует.')
                    msgBox.setWindowTitle('Ошибка')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()

        elif self.exportStack.actionListWidget.currentItem().text() == 'Экспорт стопки в *.xls (таблица)':
            if not(os.path.exists(_path) or os.path.exists(_path + '.xls')):
                _path = _path if _path[-4:] == '.xls' else _path + '.xls'
                self.exportStackToXLS(currentStack[10:], _path)
            else:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Critical)
                msgBox.setText('Файл с таким именем уже существует.')
                msgBox.setWindowTitle('Ошибка')
                msgBox.setStandardButtons(QMessageBox.Ok)
                msgBox.exec()

        elif self.exportStack.actionListWidget.currentItem().text() == 'Экспорт стопки в *.docx (карточки А7)':
            if bool(_path):
                if not(os.path.exists(_path) or os.path.exists(_path + '.docx')):
                    document = Document()
                    section = document.sections[0]
                    section.page_height = Mm(74)
                    section.page_width = Mm(105)

                    section.left_margin = Mm(10)
                    section.right_margin = Mm(10)
                    section.top_margin = Mm(10)
                    section.bottom_margin = Mm(10)

                    section.header_distance = Mm(2)
                    section.footer_distance = Mm(2)

                    myStyle = document.styles['Normal']
                    font = myStyle.font
                    font.name = 'Cambria'
                    font.size = Pt(17)

                    for card in os.listdir(currentStack):
                        if card[-5:] == '.card':
                            f = open(currentStack + '/' + card)
                            frontSide = f.readline()[:-1]
                            description = f.readline()[:-1]
                            answer = f.readline()[:-1]
                            f.close()

                            paragraph = document.add_paragraph('', style=myStyle)
                            paragraph.add_run(frontSide).bold = True
                            paragraph = document.add_paragraph('', style=myStyle)
                            paragraph.add_run(description).italic = True
                            document.add_page_break()
                            paragraph = document.add_paragraph('', style=myStyle)
                            paragraph.add_run(answer).italic = True
                            document.add_page_break()

                    _path = _path if _path[-5:] == '.docx' else _path + '.docx'
                    document.save(_path)

                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText('Стопка была успешно экспортирована в ' + _path)
                    msgBox.setWindowTitle('Информация')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()

                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Critical)
                    msgBox.setText('Файл с таким именем уже существует.')
                    msgBox.setWindowTitle('Ошибка')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()


    def getArchivePath(self):
        myFileName = QFileDialog.getOpenFileName()
        path = myFileName[0]
        self.importStack.pathQLineEdit.setText(path)


    def getPicturePath(self):
        myFileName = QFileDialog.getOpenFileName()
        self.path = myFileName[0]
        self.editor.picturePathTextBox.setText(self.path)


    def set(self):
        selectedStack = ''

        self.ui.startGamePB.clicked.connect(lambda: self.click('main', 'startGamePB', '', ''))
        self.ui.editorPB.clicked.connect(lambda: self.click('main', 'editorPB', '', ''))
        self.ui.aboutPB.clicked.connect(lambda: self.click('main', 'aboutPB', '', ''))
        self.ui.exportPB.clicked.connect(lambda: self.click('main', 'exportPB', '', ''))

        self.selectStack.backPB.clicked.connect(lambda: self.click('selectStack', 'backPB', '', ''))
        self.selectStack.startPB.clicked.connect(lambda: self.click('selectStack', 'startPB', self.selectStack.listStacks.currentItem().text() if self.selectStack.listStacks.currentItem() != None else None, ''))
        self.selectStack.newStackPB.clicked.connect(lambda: self.createStack.show())
        self.selectStack.importStackPB.clicked.connect(lambda: self.click('selectStack', 'importStackPB', '', ''))

        self.editor.editPB.clicked.connect(lambda: self.click('editor', 'editPB', self.editor.cardsList.currentItem().text() if self.editor.cardsList.currentItem() != None else None, self.editor.editingLabel.text()[15:]))
        self.editor.saveCardPB.clicked.connect(lambda: self.click('editor', 'saveCardPB', self.editor.cardsList.currentItem().text() if self.editor.cardsList.currentItem() != None else None, self.editor.editingLabel.text()[15:]))
        self.editor.goHomePB.clicked.connect(lambda: self.click('editor', 'goHomePB', '', ''))
        self.editor.browsePB.clicked.connect(lambda: self.getPicturePath())
        self.editor.rmSelCardPB.clicked.connect(lambda: self.click('editor', 'rmSelCardPB', self.editor.cardsList.currentItem().text() if self.editor.cardsList.currentItem() != None else None, self.editor.editingLabel.text()[15:]))

        self.gameWindow.submitAnswerPB.clicked.connect(lambda: self.click('game', 'sumbitAnswerPB', '', ''))
        #self.gameWindow.rollOverCardPB.clicked.connect(lambda: self.click('game', 'rollOverCardPB', self.currentCard, self.gameWindow.stackNameLabel.text()))
        self.gameWindow.knowPB.clicked.connect(lambda: self.click('game', 'knowPB', '', ''))
        self.gameWindow.dontKnowPB.clicked.connect(lambda: self.click('game', 'dontKnowPB', '', ''))
        self.gameWindow.goHomePB.clicked.connect(lambda: self.click('game', 'goHomePB', '', ''))
        self.gameWindow.rollCardPB.clicked.connect(lambda: self.click('game', 'rollCardPB', self.gameWindow.currentCardLabel.text(), self.gameWindow.stackNameLabel.text()))

        self.exportStack.goHomePB.clicked.connect(lambda: self.click('exportStack', 'goHomePB', '', ''))
        self.exportStack.exportPB.clicked.connect(lambda: self.saveArchive())

        self.createStack.createPB.clicked.connect(lambda: self.click('createStack', 'createPB', self.createStack.stackNameBox.text(), ''))

        self.importStack.browsePB.clicked.connect(lambda: self.getArchivePath())
        self.importStack.importPB.clicked.connect(lambda: self.importArchive())


    def click(self, currentForm, buttonName, var, _var):
        if currentForm == 'main':
            if buttonName == 'aboutPB':
                self.aboutWindow.show()
            elif buttonName == 'startGamePB':
                self.selectStack.listStacks.clear()
                for stack in os.listdir('user_data'):
                    self.selectStack.listStacks.addItem(stack)
                self.selectStack.label_2.setText('game')
                self.selectStack.show()
                self.ui.hide()

            elif buttonName == 'editorPB':
                self.selectStack.listStacks.clear()
                for stack in os.listdir('user_data'):
                    self.selectStack.listStacks.addItem(stack)
                self.selectStack.label_2.setText('editor')
                self.selectStack.show()
                self.ui.hide()

            elif buttonName == 'exportPB':
                self.selectStack.listStacks.clear()
                for stack in os.listdir('user_data'):
                    self.selectStack.listStacks.addItem(stack)
                self.selectStack.label_2.setText('export')
                self.selectStack.show()
                self.ui.hide()

        elif currentForm == 'selectStack':
            if buttonName == 'backPB':
                self.selectStack.listStacks.clear()
                self.selectStack.hide()
                self.ui.show()

            elif buttonName == 'startPB':
                if var != None:
                    if self.selectStack.label_2.text() == 'game':
                        allcards = 0
                        for b in os.listdir('user_data/' + var):
                            if b[-5:] == '.card':
                                allcards += 1

                        if allcards > 0:
                            self.selectStack.listStacks.clear()
                            self.selectStack.hide()

                            self.gameWindow.stackNameLabel.setText(var)
                            self.gameWindow.show()

                            currentStack = self.gameWindow.stackNameLabel.text()
                            a = open('tmp/currentStack.tmp', 'w').write(currentStack)

                            self.gameWindow.myProgressBar.setValue(0)
                            self.gameWindow.showResultLabel.setText('OK/BAD/END')
                            self.gameWindow.showResultLabel.setStyleSheet('')
                            self.gameWindow.countProgressLabel.setText('0/x')
                            self.gameWindow.nextBoolLabel.setText('')
                            self.gameWindow.backBoolLabel.setText('')
                            self.gameWindow.acceptBoolLabel.setText('')
                            self.gameWindow.knowBoolLabel.setText('')
                            self.gameWindow.dontKnowBoolLabel.setText('')
                            self.gameWindow.rollBoolLabel.setText('')
                            self.gameWindow.progressLabel2.setText('0')

                            self.startGameThread_instance = gameThreadClass(gameWindow=self.gameWindow)
                            self.startGameThread_instance.start()


                            self.setPValue_instance = setProgressBarValueClass(gameWindow=self.gameWindow)
                            self.setPValue_instance.start()
                        else:
                            msgBox = QMessageBox()
                            msgBox.setIcon(QMessageBox.Critical)
                            msgBox.setText('Выбранная стопка пуста. Нажмите ОК для того, чтобы перейти в редактор карточек.')
                            msgBox.setWindowTitle('Ошибка')
                            msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
                            userReply = msgBox.exec()

                            if userReply == QMessageBox.Ok:
                                self.selectStack.hide()
                                self.editor.editingLabel.setText('Редактируется: ' + var)
                                self.editor.show()


                    elif self.selectStack.label_2.text() == 'editor':
                        for card in os.listdir('user_data/' + var):
                            if card[-5:] == '.card':
                                self.editor.cardsList.addItem(card[:-5])
                        self.selectStack.hide()
                        self.editor.editingLabel.setText('Редактируется: ' + var)
                        self.editor.show()

                    else:
                        self.exportStack.stackLabel.setText('user_data/' + var)
                        self.exportStack.actionListWidget.clear()
                        self.exportStack.actionListWidget.addItem('Экспорт стопки в архив (для копирования)')
                        self.exportStack.actionListWidget.addItem('Экспорт стопки в *.xls (таблица)')
                        self.exportStack.actionListWidget.addItem('Экспорт стопки в *.docx (карточки А7)')
                        self.exportStack.show()
                        self.selectStack.hide()

                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Critical)
                    msgBox.setText('Вы не выбрали стопку.')
                    msgBox.setWindowTitle('Ошибка')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()

            elif buttonName == 'importStackPB':
                self.importStack.show()


        elif currentForm == 'editor':
            if buttonName == 'editPB':
                if var != None:
                    self.editor.frontSideLineEditBox.clear()
                    self.editor.descTextEditBox.clear()
                    self.editor.backSideTextEditBox.clear()
                    self.editor.picturePathTextBox.clear()
                    self.editor.cardFileNameBox.clear()

                    streamReader = open('user_data/' + _var + '/' + var + '.card')
                    cardCaption = streamReader.readline()[:-1]
                    description = streamReader.readline()[:-1]
                    backSideCrd = streamReader.readline()[:-1]
                    picturePath = streamReader.readline()
                    streamReader.close()

                    self.editor.frontSideLineEditBox.setText(cardCaption)
                    self.editor.descTextEditBox.clear()
                    self.editor.descTextEditBox.insertPlainText(description)
                    self.editor.backSideTextEditBox.setText(backSideCrd)
                    self.editor.picturePathTextBox.setText(picturePath)
                    self.editor.cardFileNameBox.setText(var)

                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Critical)
                    msgBox.setText('Вы не выбрали карточку.')
                    msgBox.setWindowTitle('Ошибка')
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    msgBox.exec()

            elif buttonName == 'saveCardPB':
                if var != None:
                    os.remove('user_data/' + _var + '/' + var + '.card')
                    streamWriter = open('user_data/' + _var + '/' + self.editor.cardFileNameBox.text() + '.card', 'w')
                else:
                    if not(os.path.exists('user_data/' + _var + '/' + self.editor.cardFileNameBox.text() + '.card')):
                        streamWriter = open('user_data/' + _var + '/' + self.editor.cardFileNameBox.text() + '.card', 'w')
                try:
                    if self.editor.cardFileNameBox.text() + '.card' != '.card':
                        streamWriter.write(self.editor.frontSideLineEditBox.text() + '\n')
                        streamWriter.write(self.editor.descTextEditBox.toPlainText() + '\n')
                        streamWriter.write(self.editor.backSideTextEditBox.text() + '\n')
                        if len(self.editor.picturePathTextBox.text()) != 0:
                            shutil.copyfile(self.editor.picturePathTextBox.text(), 'user_data/' + _var + '/' + self.editor.cardFileNameBox.text() + '.' + self.editor.picturePathTextBox.text()[-3:])
                        streamWriter.write('//nopicture//' if len(self.editor.picturePathTextBox.text()) == 0 else 'user_data/' + _var + '/' + self.editor.cardFileNameBox.text() + '.' + self.editor.picturePathTextBox.text()[-3:])

                        streamWriter.close()
                    else:
                        if os.path.exists('user_data/' + _var + '/.card'):
                            os.remove('user_data/' + _var + '/.card')
                        msgBox = QMessageBox()
                        msgBox.setIcon(QMessageBox.Critical)
                        msgBox.setText('Следует указать имя файла.')
                        msgBox.setWindowTitle('Ошибка')
                        msgBox.setStandardButtons(QMessageBox.Ok)
                        msgBox.exec()
                except:
                    pass


                self.editor.cardsList.clear()
                for card in os.listdir('user_data/' + _var):
                    if card[-5:] == '.card':
                        self.editor.cardsList.addItem(card[:-5])

                self.editor.frontSideLineEditBox.clear()
                self.editor.descTextEditBox.clear()
                self.editor.backSideTextEditBox.clear()
                self.editor.picturePathTextBox.clear()
                self.editor.cardFileNameBox.clear()

            elif buttonName == 'goHomePB':
                self.editor.cardsList.clear()
                self.selectStack.listStacks.clear()
                self.ui.show()
                self.editor.hide()

                self.editor.frontSideLineEditBox.clear()
                self.editor.descTextEditBox.clear()
                self.editor.backSideTextEditBox.clear()
                self.editor.picturePathTextBox.clear()
                self.editor.cardFileNameBox.clear()

            elif buttonName == 'rmSelCardPB':
                if var != None:
                    os.remove('user_data/' + _var + '/' + var + '.card')
                else:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Question)
                    msgBox.setText('Вы уверены, что хотите удалить стопку? (Действие нельзя отменить)')
                    msgBox.setWindowTitle('Удаление стопки')
                    msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                    userReply = msgBox.exec()
                    if userReply == QMessageBox.Yes:
                        shutil.rmtree('user_data/' + _var)
                        msgBox = QMessageBox()
                        msgBox.setIcon(QMessageBox.Information)
                        msgBox.setText('Стопка успешно удалена. Запустите программу снова для обновления информации')
                        msgBox.setWindowTitle('Удаление стопки')
                        msgBox.setStandardButtons(QMessageBox.Ok)
                        msgBox.exec()
                        exit()


                self.editor.cardsList.clear()
                for card in os.listdir('user_data/' + _var):
                    if card[-5:] == '.card':
                        self.editor.cardsList.addItem(card[:-5])

                self.editor.frontSideLineEditBox.clear()
                self.editor.descTextEditBox.clear()
                self.editor.backSideTextEditBox.clear()
                self.editor.picturePathTextBox.clear()
                self.editor.cardFileNameBox.clear()


        elif currentForm == 'game':
            if buttonName == 'sumbitAnswerPB':
                self.gameWindow.acceptBoolLabel.setText('True')

            elif buttonName == 'rollCardPB':
                f = open('user_data/' + _var + '/' + var)
                frontSide = f.readline()[:-1]
                description = f.readline()[:-1]
                answer = f.readline()[:-1]
                f.close()
                if not(bool(self.gameWindow.rollBoolLabel.text())):
                    self.gameWindow.rollBoolLabel.setText('True')
                    self.gameWindow.cardDescriptionTextEdit.clear()
                    self.gameWindow.cardDescriptionTextEdit.insertPlainText(answer)
                    self.gameWindow.update()
                else:
                    self.gameWindow.rollBoolLabel.setText('')
                    self.gameWindow.cardDescriptionTextEdit.clear()
                    self.gameWindow.cardDescriptionTextEdit.insertPlainText(description)
                    self.gameWindow.update()
            elif buttonName == 'knowPB':
                self.gameWindow.knowBoolLabel.setText('True')
            elif buttonName == 'dontKnowPB':
                self.gameWindow.dontKnowBoolLabel.setText('True')
            elif buttonName == 'goHomePB':
                _correctCount = int(self.gameWindow.correctCount.text())
                _allcards = int(self.gameWindow.countProgressLabel.text().split('/')[1])
                self.gameWindow.showResultLabel.setText('END')
                self.gameWindow.showResultLabel.setStyleSheet('color: blue;')
                endResult = 'Последний результат: ' + str(_correctCount) + '/' + str(_allcards)
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(endResult)
                msgBox.setWindowTitle('Информация')
                msgBox.setStandardButtons(QMessageBox.Ok)
                msgBox.exec()
                sys.exit()


        elif currentForm == 'exportStack':
            if buttonName == 'goHomePB':
                self.ui.show()
                self.exportStack.hide()

        elif currentForm == 'createStack':
            if buttonName == 'createPB':
                if len(var) != 0 and var != ' ':
                    if not(os.path.exists('user_data/' + var)):
                        os.mkdir('user_data/' + var)
                        self.createStack.hide()
                        self.selectStack.hide()
                        for card in os.listdir('user_data/' + var):
                            if card[-5:] == '.card':
                                self.editor.cardsList.addItem(card[:-5])


                        msgBox = QMessageBox()
                        msgBox.setIcon(QMessageBox.Information)
                        msgBox.setText('Стопка карт была успешно создана.')
                        msgBox.setWindowTitle('Информация')
                        msgBox.setStandardButtons(QMessageBox.Ok)
                        msgBox.exec()

                        self.editor.editingLabel.setText('Редактируется: ' + var)
                        self.editor.show()
                    else:
                        msgBox = QMessageBox()
                        msgBox.setIcon(QMessageBox.Critical)
                        msgBox.setText('Стопка с таким именем уже существует.')
                        msgBox.setWindowTitle('Ошибка')
                        msgBox.setStandardButtons(QMessageBox.Ok)
                        msgBox.exec()



if __name__ == '__main__':
    if not(os.path.exists('user_data')):
        os.chdir(__file__[:-8])

    app = QApplication(sys.argv)
    ex = App()
    app.exec_()
